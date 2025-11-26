from __future__ import annotations
import io
import json
import hashlib
from datetime import datetime
from typing import Dict, Any, List, Tuple

from supabase import create_client, Client
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ----------------------------
# Config
# ----------------------------
BUCKET_NAME = "nhcma-booklets"
TABLE_NAME  = "submissions"

# Optional: standard section ordering for nicer layout (keys must match payload_json fields)
ORG_SECTION_KEYS = [
    ("Applicant & Organization", [
        "applicant_name", "email", "phone", "org_name", "mission",
        "exec_name", "exec_email", "exec_phone"
    ]),
    ("Eligibility", [
        "eligibility.nonprofit", "eligibility.benefit_gnh", "eligibility.report_at_winter_meeting_2026"
    ]),
    ("Project", [
        "project_title", "q1_issue", "q2_align", "q3_benefit", "description",
        "budget_text", "budget_total", "timeline", "evaluation"
    ]),
    ("Attachments", [
        "proposal", "budget", "other"
    ]),
]

STUDENT_SECTION_KEYS = [
    ("Applicant", [
        "applicant_name", "email", "phone", "school", "grad_date"
    ]),
    ("Advisor (optional)", [
        "advisor_name", "advisor_title", "advisor_email"
    ]),
    ("Eligibility", [
        "eligibility.enrolled_qu_yale", "eligibility.report_at_winter_meeting_2026"
    ]),
    ("Project", [
        "project_title", "q1_issue", "q2_align", "q3_benefit", "description",
        "budget_text", "budget_total", "timeline", "evaluation"
    ]),
    ("Attachments", [
        "proposal", "budget", "other", "cv", "support_letter"
    ]),
]

# ----------------------------
# Supabase helpers
# ----------------------------

def get_supabase(url: str, key: str) -> Client:
    return create_client(url, key)


def list_submissions(sb: Client, where: Dict[str, Any] | None = None) -> List[Dict[str, Any]]:
    q = sb.table(TABLE_NAME).select("*")
    # only apply filters if explicitly requested and column likely exists
    if where and isinstance(where, dict):
        try:
            for k, v in where.items():
                q = q.eq(k, v)
        except Exception:
            # if the column doesn't exist, just skip filter
            pass
    res = q.execute()
    return res.data or []


def update_submission_paths(sb: Client, submission_id: Any, path_docx: str, version: int | None = None) -> None:
    payload = {
        "booklet_docx_path": path_docx,
        "booklet_updated_at": datetime.utcnow().isoformat() + "Z",
    }
    if version is not None:
        payload["booklet_version"] = version
    sb.table(TABLE_NAME).update(payload).eq("id", submission_id).execute()


def upload_bytes(sb: Client, bucket: str, path: str, data: bytes, content_type: str) -> None:
    storage = sb.storage.from_(bucket)
    # supabase-py expects "contentType" (camelCase) and upsert as a string
    resp = storage.upload(
        path=path,
        file=data,
        file_options={"contentType": content_type, "upsert": "true"},
    )
    # Basic sanity check: raise if error-like payload returned
    if isinstance(resp, dict) and resp.get("error"):
        raise RuntimeError(f"Storage upload failed: {resp['error']}")


def make_signed_url(sb: Client, bucket: str, path: str, expires_in_seconds: int = 86400) -> str:
    storage = sb.storage.from_(bucket)
    res = storage.create_signed_url(path, expires_in_seconds)
    return res.get("signedURL") or res.get("signed_url") or ""

# ----------------------------
# Payload utilities
# ----------------------------

def _ensure_dict(payload_json: Any) -> Dict[str, Any]:
    if isinstance(payload_json, dict):
        return payload_json
    if isinstance(payload_json, str):
        return json.loads(payload_json)
    raise TypeError("payload_json is neither dict nor str")


def payload_hash(payload: Dict[str, Any]) -> str:
    # Stable hash for change detection, if you ever want it
    canonical = json.dumps(payload, sort_keys=True, ensure_ascii=False)
    return hashlib.sha256(canonical.encode("utf-8")).hexdigest()


def get_value(payload: Dict[str, Any], dotted_key: str):
    # Supports keys like "eligibility.nonprofit"
    cur = payload
    for part in dotted_key.split("."):
        if isinstance(cur, dict) and part in cur:
            cur = cur[part]
        else:
            return None
    return cur


def fmt_bool(v) -> str:
    if isinstance(v, bool):
        return "✅ Yes" if v else "❌ No"
    # Some forms store "yes"/"no" strings
    if isinstance(v, str):
        if v.strip().lower() in {"yes", "true", "y"}:
            return "✅ Yes"
        if v.strip().lower() in {"no", "false", "n"}:
            return "❌ No"
    return str(v) if v is not None else "—"


def human_label(key: str) -> str:
    # Turn snake_case or dotted into Title Case labels
    parts = key.split(".")[-1].replace("_", " ")
    return parts.strip().capitalize()

# ----------------------------
# DOCX builder
# ----------------------------

def build_booklet_docx(submission_row: Dict[str, Any]) -> bytes:
    """Create a DOCX booklet from a submissions row. Returns bytes."""
    payload = _ensure_dict(submission_row.get("payload_json"))
    track   = submission_row.get("track", "") or payload.get("track", "")
    sub_id  = submission_row.get("id")
    created = submission_row.get("created_at")

    doc = Document()

    # Cover
    title = payload.get("project_title") or payload.get("title") or "Grant Submission"
    h = doc.add_heading(level=0)
    run = h.add_run(f"NHCMA Foundation — {str(track).title()} Track")
    run.bold = True
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    h2 = doc.add_heading(level=1)
    h2.add_run(title)
    h2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"Submission ID: {sub_id}  •  Submitted: {created}")

    doc.add_page_break()

    # Decide section template by track
    sections = ORG_SECTION_KEYS if str(track).lower().startswith("org") else STUDENT_SECTION_KEYS

    # Render sections
    for section_title, keys in sections:
        doc.add_heading(section_title, level=2)
        for key in keys:
            val = get_value(payload, key)
            label = human_label(key)
            if isinstance(val, bool) or (isinstance(val, str) and val.lower() in {"yes", "no", "true", "false"}):
                pretty = fmt_bool(val)
            else:
                pretty = str(val) if val not in (None, "") else "—"

            # URLs: make it explicit in text (python-docx has no real hyperlink for external links without XML)
            if key in {"proposal", "budget", "other", "cv", "support_letter"} and pretty not in {"—", "None"}:
                doc.add_paragraph(f"{label}: {pretty}")
            else:
                para = doc.add_paragraph()
                run = para.add_run(f"{label}: ")
                run.bold = True
                para.add_run(pretty)
        doc.add_paragraph("")

    # --- Attachments Section (print clickable URLs) ---
    uploads = submission_row.get("uploads_json") or {}
    if isinstance(uploads, str):
        try:
            uploads = json.loads(uploads)
        except Exception:
            uploads = {}

    doc.add_heading("Attachments", level=2)
    ATTACHMENT_KEYS = ["proposal", "budget", "other", "cv", "support_letter"]
    any_written = False
    for key in ATTACHMENT_KEYS:
        # Skip CV/support letter for org track
        if str(track).lower().startswith("org") and key in ("cv", "support_letter"):
            continue

        # Prefer uploads_json → fallback to payload
        url = None
        v = uploads.get(key)
        if isinstance(v, str) and v.strip():
            url = v
        else:
            pv = get_value(payload, key) or payload.get(key)
            if isinstance(pv, str) and pv.strip():
                url = pv

        para = doc.add_paragraph()
        run = para.add_run(f"{human_label(key)}: ")
        run.bold = True
        para.add_run(url if url else "—")
        any_written = any_written or bool(url)

    if not any_written:
        doc.add_paragraph("No attachments uploaded.")

    # Optional: dump any extra fields not covered above under an Appendix
    covered = {k for _, lst in sections for k in lst}
    extra_keys = sorted(set(payload.keys()) - {k.split(".")[0] for k in covered})
    if extra_keys:
        doc.add_page_break()
        doc.add_heading("Appendix — Additional Fields", level=2)
        for k in extra_keys:
            v = payload.get(k)
            if isinstance(v, (dict, list)):
                try:
                    v = json.dumps(v, ensure_ascii=False)
                except Exception:
                    v = str(v)
            para = doc.add_paragraph()
            run = para.add_run(f"{human_label(k)}: ")
            run.bold = True
            para.add_run(str(v) if v not in (None, "") else "—")

    # Typography tweaks
    for p in doc.paragraphs:
        for run in p.runs:
            run.font.size = Pt(11)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

# ----------------------------
# Batch builder
# ----------------------------

def build_and_store_docx_for_row(sb: Client, row: Dict[str, Any], version: int | None = None) -> Tuple[str, str]:
    """Builds a DOCX, uploads to Storage, updates the row. Returns (path, signed_url)."""
    docx_bytes = build_booklet_docx(row)

    track   = (row.get("track") or "").replace(" ", "")
    sub_id  = row.get("id")
    ts      = datetime.utcnow().strftime("%Y%m%d-%H%M%S")

    path = f"booklets/{sub_id}/NHCMA_Grants_{track}_{sub_id}_{ts}.docx"
    upload_bytes(sb, BUCKET_NAME, path, docx_bytes,
                 content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    update_submission_paths(sb, sub_id, path_docx=path, version=version)

    # Short-lived signed URL (e.g., 24h). For judge UI, consider generating on demand instead.
    signed = make_signed_url(sb, BUCKET_NAME, path, expires_in_seconds=24*3600)
    return path, signed


def build_all_booklets(sb: Client, where: Dict[str, Any] | None = None, start_version: int = 1) -> List[Dict[str, Any]]:
    """Iterate submissions and build booklets. Returns a report list per row."""
    rows = list_submissions(sb, where=where)
    report: List[Dict[str, Any]] = []
    version = start_version
    for row in rows:
        try:
            path, signed = build_and_store_docx_for_row(sb, row, version=version)
            report.append({
                "id": row.get("id"),
                "track": row.get("track"),
                "docx_path": path,
                "signed_url": signed,
                "status": "ok",
            })
            version += 1
        except Exception as e:
            report.append({
                "id": row.get("id"),
                "track": row.get("track"),
                "error": str(e),
                "status": "error",
            })
    return report

# ----------------------------
# (Optional) Streamlit admin hook
# ----------------------------
"""
Example wiring in your admin page:

import streamlit as st
from nhcma_booklet_builder import get_supabase, build_all_booklets

sb = get_supabase(st.secrets["SUPABASE_URL"], st.secrets["SUPABASE_SERVICE_ROLE_KEY"])

st.header("Build Booklets for Judging")
if st.button("Build All Booklets Now"):
    with st.spinner("Building…"):
        report = build_all_booklets(sb, where={"status": "submitted"})
    st.success("Done.")
    st.dataframe(report)

In the judge view (to show a fresh signed URL from stored path):

path = row.get("booklet_docx_path")
if path:
    url = make_signed_url(sb, BUCKET_NAME, path, expires_in_seconds=24*3600)
    st.link_button("Download Booklet (DOCX)", url)
else:
    st.info("No booklet yet — ask admin to build.")
"""
