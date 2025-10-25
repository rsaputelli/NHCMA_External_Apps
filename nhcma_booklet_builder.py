from __future__ import annotations

import io
import json
import hashlib
from datetime import datetime
from typing import Dict, Any, List, Tuple, Optional  # <-- include Optional

from supabase import create_client, Client
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from urllib.parse import quote

# --- Edge Function Config ---
PROJECT_REF = "icjunpjliexaacexjgwy"  # your Supabase project ref
EDGE_BASE   = f"https://{PROJECT_REF}.supabase.co/functions/v1"

def edge_signed_download_url(bucket: str, path: str) -> str:
    return f"{EDGE_BASE}/signed-download?bucket={quote(bucket)}&path={quote(path, safe='/')}"

def _bucket_and_key_from_url(u: str) -> tuple[str, str]:
    if isinstance(u, str) and u.startswith("http"):
        m = "/storage/v1/object/public/"
        if m in u:
            rest = u.split(m, 1)[1].split("?", 1)[0]
            bucket, _, key = rest.partition("/")
            return bucket, key
        m = "/storage/v1/object/sign/"
        if m in u:
            rest = u.split(m, 1)[1].split("?", 1)[0]
            bucket, _, key = rest.partition("/")
            return bucket, key
    return "nhcma-uploads", (u or "").lstrip("/")

def _ensure_dict(payload_json):
    if isinstance(payload_json, dict):
        return payload_json
    if isinstance(payload_json, str):
        return json.loads(payload_json)
    return {}

def _normalized_track(row: dict) -> str:
    t = (row.get("track") or "").strip().lower()
    if not t:
        payload = _ensure_dict(row.get("payload_json"))
        t = (payload.get("track") or "").strip().lower()
    if "stud" in t: return "student"
    if "org" in t or "organization" in t: return "organization"
    return "student"

# ----------------------------
# Config
# ----------------------------
BUCKET_NAME = "nhcma-booklets"
TABLE_NAME  = "submissions"

# Attachment key set we will use consistently
ATTACHMENT_KEYS = ["proposal", "budget", "other", "cv", "support_letter"]
ATTACHMENT_SET = set(ATTACHMENT_KEYS)

# Optional: standard section ordering for nicer layout (keys must match payload_json fields)
# NOTE: we include "Attachments" here for organization of labels, BUT we'll SKIP these keys
# during the section rendering so that only the dedicated Attachments section renders links.
ORG_SECTION_KEYS = [
    ("Applicant & Organization", [
        "applicant_name", "email", "phone", "org_name", "mission",
        "exec_name", "exec_email", "exec_phone"
    ]),
    ("Eligibility", [
        "eligibility.nonprofit", "eligibility.benefit_gnh", "eligibility.report_at_winter_meeting_2026"
    ]),
    ("Project", [
        "project_title", "q1_issue", "q2_align", "q3_benefit", "description",
        "budget_text", "budget_total", "timeline", "evaluation"
    ]),
    # Keys listed, but they'll be skipped here to avoid duplicate rendering
    ("Attachments", [
        "proposal", "budget", "other"
    ]),
]

STUDENT_SECTION_KEYS = [
    ("Applicant", [
        "applicant_name", "email", "phone", "school", "grad_date"
    ]),
    ("Advisor (optional)", [
        "advisor_name", "advisor_title", "advisor_email"
    ]),
    ("Eligibility", [
        "eligibility.enrolled_qu_yale", "eligibility.report_at_winter_meeting_2026"
    ]),
    ("Project", [
        "project_title", "q1_issue", "q2_align", "q3_benefit", "description",
        "budget_text", "budget_total", "timeline", "evaluation"
    ]),
    # Keys listed, but they'll be skipped here to avoid duplicate rendering
    ("Attachments", [
        "proposal", "budget", "other", "cv", "support_letter"
    ]),
]

# ----------------------------
# Supabase helpers
# ----------------------------

def get_supabase(url: str, key: str) -> Client:
    return create_client(url, key)


def list_submissions(sb: Client, where: Dict[str, Any] | None = None) -> List[Dict[str, Any]]:
    q = sb.table(TABLE_NAME).select("*")
    # only apply filters if explicitly requested and column likely exists
    if where and isinstance(where, dict):
        try:
            for k, v in where.items():
                q = q.eq(k, v)
        except Exception:
            # if the column doesn't exist, just skip filter
            pass
    res = q.execute()
    return res.data or []


def update_submission_paths(sb: Client, submission_id: Any, path_docx: str, version: int | None = None) -> None:
    payload = {
        "booklet_docx_path": path_docx,
        "booklet_updated_at": datetime.utcnow().isoformat() + "Z",
    }
    if version is not None:
        payload["booklet_version"] = version
    sb.table(TABLE_NAME).update(payload).eq("id", submission_id).execute()


def upload_bytes(sb: Client, bucket: str, path: str, data: bytes, content_type: str) -> None:
    storage = sb.storage.from_(bucket)
    # supabase-py expects "contentType" (camelCase) and upsert as a string
    resp = storage.upload(
        path=path,
        file=data,
        file_options={"contentType": content_type, "upsert": "true"},
    )
    # Basic sanity check: raise if error-like payload returned
    if isinstance(resp, dict) and resp.get("error"):
        raise RuntimeError(f"Storage upload failed: {resp['error']}")


def make_signed_url(sb: Client, bucket: str, path: str, expires_in_seconds: int = 86400) -> str:
    storage = sb.storage.from_(bucket)
    res = storage.create_signed_url(path, expires_in_seconds)
    return res.get("signedURL") or res.get("signed_url") or ""

# ----------------------------
# Payload utilities
# ----------------------------

def _ensure_dict(payload_json: Any) -> Dict[str, Any]:
    if isinstance(payload_json, dict):
        return payload_json
    if isinstance(payload_json, str):
        return json.loads(payload_json)
    raise TypeError("payload_json is neither dict nor str")


def payload_hash(payload: Dict[str, Any]) -> str:
    # Stable hash for change detection, if you ever want it
    canonical = json.dumps(payload, sort_keys=True, ensure_ascii=False)
    return hashlib.sha256(canonical.encode("utf-8")).hexdigest()


def get_value(payload: Dict[str, Any], dotted_key: str):
    # Supports keys like "eligibility.nonprofit"
    cur = payload
    for part in dotted_key.split("."):
        if isinstance(cur, dict) and part in cur:
            cur = cur[part]
        else:
            return None
    return cur


def fmt_bool(v) -> str:
    if isinstance(v, bool):
        return "✅ Yes" if v else "❌ No"
    # Some forms store "yes"/"no" strings
    if isinstance(v, str):
        if v.strip().lower() in {"yes", "true", "y"}:
            return "✅ Yes"
        if v.strip().lower() in {"no", "false", "n"}:
            return "❌ No"
    return str(v) if v is not None else "—"


def human_label(key: str) -> str:
    # Turn snake_case or dotted into Title Case labels
    parts = key.split(".")[-1].replace("_", " ")
    return parts.strip().capitalize()

# ----------------------------
# DOCX helpers (clickable links)
# ----------------------------

def _add_hyperlink(paragraph, text: str, url: str):
    """
    Insert a clickable hyperlink using a simple field code:
      <w:fldSimple w:instr='HYPERLINK "url"'>...</w:fldSimple>
    This renders as a blue, underlined clickable link in Word.
    """
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), f'HYPERLINK "{url}"')

    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0000FF")
    rPr.append(color)

    r.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    r.append(t)

    fld.append(r)
    paragraph._p.append(fld)
    return paragraph

# ----------------------------
# Track helper
# ----------------------------

def _normalized_track(row: dict) -> str:
    """
    Prefer row['track'], fall back to payload_json['track'].
    Returns 'student' or 'organization' (default 'student' for safety).
    """
    t = (row.get("track") or "").strip().lower()
    if not t:
        try:
            payload = _ensure_dict(row.get("payload_json"))
            t = (payload.get("track") or "").strip().lower()
        except Exception:
            t = ""
    if "stud" in t:
        return "student"
    if "org" in t or "organization" in t:
        return "organization"
    return "student"

# ----------------------------
# DOCX builder
# ----------------------------

def build_booklet_docx(submission_row: Dict[str, Any]) -> bytes:
    """Create a DOCX booklet from a submissions row. Returns bytes."""
    payload = _ensure_dict(submission_row.get("payload_json"))
    sub_id  = submission_row.get("id")
    created = submission_row.get("created_at")
    track_norm = _normalized_track(submission_row)

    doc = Document()

    # Cover
    title = payload.get("project_title") or payload.get("title") or "Grant Submission"
    h = doc.add_heading(level=0)
    run = h.add_run(f"NHCMA Foundation — {track_norm.title()} Track")
    run.bold = True
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    h2 = doc.add_heading(level=1)
    h2.add_run(title)
    h2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"Submission ID: {sub_id}  •  Submitted: {created}")

    doc.add_page_break()

    # Decide section template by track (robust)
    sections = STUDENT_SECTION_KEYS if track_norm == "student" else ORG_SECTION_KEYS

    # If any advisor fields present, always show an Advisor section after first section
    if any((payload.get("advisor_name"), payload.get("advisor_title"), payload.get("advisor_email"))):
        sections = list(sections)  # copy if tuple
        # Only insert Advisor section if not already included by the chosen template
        advisor_keys = ("advisor_name", "advisor_title", "advisor_email")
        if not any("Advisor" in title for title, _ in sections):
            sections.insert(1, ("Advisor (optional)", list(advisor_keys)))

    # Render sections (SKIP attachment keys here to avoid duplicates)
    for section_title, keys in sections:
        doc.add_heading(section_title, level=2)
        for key in keys:
            if key in ATTACHMENT_SET:
                # attachments are rendered only in the dedicated section below
                continue

            val = get_value(payload, key)
            label = human_label(key)
            if isinstance(val, bool) or (isinstance(val, str) and val.lower() in {"yes", "no", "true", "false"}):
                pretty = fmt_bool(val)
            else:
                pretty = str(val) if val not in (None, "") else "—"

            para = doc.add_paragraph()
            run = para.add_run(f"{label}: ")
            run.bold = True
            para.add_run(pretty)
        doc.add_paragraph("")

    # --- Attachments Section (clickable URLs) ---
    uploads = submission_row.get("uploads_json") or {}
    if isinstance(uploads, str):
        try:
            uploads = json.loads(uploads)
        except Exception:
            uploads = {}

    doc.add_heading("Attachments", level=2)
    any_written = False

    for key in ATTACHMENT_KEYS:
        # NOTE: removed the org-track skip so Org also shows CV/Support Letter when present

        # Prefer uploads_json → fallback to payload
        url = None
        v = uploads.get(key)
        if isinstance(v, str) and v.strip():
            url = v
        else:
            pv = get_value(payload, key) or payload.get(key)
            if isinstance(pv, str) and pv.strip():
                url = pv

        para = doc.add_paragraph()
        run = para.add_run(f"{human_label(key)}: ")
        run.bold = True

        if url:
            bucket, obj_key = _bucket_and_key_from_url(url)
            signed_edge_url = edge_signed_download_url(bucket, obj_key)
            _add_hyperlink(para, url, signed_edge_url)   # clickable 🔗 via Edge Function
            any_written = True
        else:
            para.add_run("—")

    if not any_written:
        doc.add_paragraph("No attachments uploaded.")

    # Optional: dump any extra fields not covered above under an Appendix
    used_keys = {k for _, lst in sections for k in lst} | ATTACHMENT_SET
    extra_keys = [k for k in payload.keys() if k not in used_keys]
    if extra_keys:
        doc.add_page_break()
        doc.add_heading("Appendix — Additional Fields", level=2)
        for k in extra_keys:
            v = payload.get(k)
            if isinstance(v, (dict, list)):
                try:
                    v = json.dumps(v, ensure_ascii=False)
                except Exception:
                    v = str(v)
            para = doc.add_paragraph()
            run = para.add_run(f"{human_label(k)}: ")
            run.bold = True
            para.add_run(str(v) if v not in (None, "") else "—")

    # Typography tweaks
    for p in doc.paragraphs:
        for run in p.runs:
            run.font.size = Pt(11)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

# ----------------------------
# Batch builder
# ----------------------------

def build_and_store_docx_for_row(sb: Client, row: Dict[str, Any], version: int | None = None) -> Tuple[str, str]:
    """Builds a DOCX, uploads to Storage, updates the row. Returns (path, signed_url)."""
    docx_bytes = build_booklet_docx(row)

    track   = (_normalized_track(row) or "").replace(" ", "")
    sub_id  = row.get("id")
    ts      = datetime.utcnow().strftime("%Y%m%d-%H%M%S")

    path = f"booklets/{sub_id}/NHCMA_Grants_{track}_{sub_id}_{ts}.docx"
    upload_bytes(sb, BUCKET_NAME, path, docx_bytes,
                 content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    update_submission_paths(sb, sub_id, path_docx=path, version=version)

    # Signed URL for the booklet (24h). For judge UI, you can regenerate on demand.
    signed = make_signed_url(sb, BUCKET_NAME, path, expires_in_seconds=24*3600)
    return path, signed


def build_all_booklets(
    sb: Client,
    where: Optional[Dict[str, Any]] = None,
    start_version: Optional[int] = None,
):
    """Iterate submissions and build booklets. Returns a report list per row."""
    rows = list_submissions(sb, where=where)
    report: List[Dict[str, Any]] = []
    version = start_version if start_version is not None else 1  # harden

    for row in rows:
        try:
            path, signed = build_and_store_docx_for_row(sb, row, version=version)
            report.append({
                "id": row.get("id"),
                "track": _normalized_track(row),
                "docx_path": path,
                "signed_url": signed,
                "status": "ok",
            })
            version += 1
        except Exception as e:
            report.append({
                "id": row.get("id"),
                "track": row.get("track"),
                "error": str(e),
                "status": "error",
            })
    return report

# ----------------------------
# (Optional) Streamlit admin hook
# ----------------------------
"""
Example wiring in your admin page:

import streamlit as st
from nhcma_booklet_builder import get_supabase, build_all_booklets

sb = get_supabase(st.secrets["SUPABASE_URL"], st.secrets["SUPABASE_SERVICE_ROLE_KEY"])

st.header("Build Booklets for Judging")
if st.button("Build All Booklets Now"):
    with st.spinner("Building…"):
        report = build_all_booklets(sb, where={"status": "submitted"})
    st.success("Done.")
    st.dataframe(report)

In the judge view (to show a fresh signed URL from stored path):

path = row.get("booklet_docx_path")
if path:
    url = make_signed_url(sb, BUCKET_NAME, path, expires_in_seconds=24*3600)
    st.link_button("Download Booklet (DOCX)", url)
else:
    st.info("No booklet yet — ask admin to build.")
"""
